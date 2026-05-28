from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "NetPro_Buku_Panduan_Guidebook.docx"
LOGO = ROOT / "game" / "logo" / "NetPro-Logo.png"


COLORS = {
    "navy": "0D1117",
    "panel": "0F1923",
    "blue": "4060C0",
    "royal": "5070C0",
    "cyan": "60C0F0",
    "pale": "F0F7FF",
    "soft": "EAF5FF",
    "success": "76FF03",
    "warning": "FFD740",
    "ink": "182033",
    "muted": "5D6B82",
    "border": "B8D8FF",
    "brown": "604050",
    "white": "FFFFFF",
}


def rgb(hex_value):
    value = hex_value.replace("#", "")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8D8FF", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=COLORS["muted"])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["blue"])
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color, before, after in [
        ("Heading 1", 17, COLORS["blue"], 12, 8),
        ("Heading 2", 13.5, COLORS["royal"], 8, 5),
        ("Heading 3", 11.5, COLORS["panel"], 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15


def set_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "NetPro: Magang Jaringan"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_run_font(run, size=9, color=COLORS["muted"], bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Buku Panduan Game | Halaman "
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        set_run_font(run, size=9, color=COLORS["muted"])
    add_page_field(fp)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_run_font(r, size=10, color=COLORS["cyan"], bold=True)
    return p


def add_title(doc, text, subtitle=None):
    add_kicker(doc, "Buku Panduan")
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(14)
        r = sp.add_run(subtitle)
        set_run_font(r, size=12, color=COLORS["muted"], italic=True)


def add_section_title(doc, number, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number:02d}  {title}")
    set_run_font(r, size=18, color=COLORS["blue"], bold=True)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        sr = sp.add_run(subtitle)
        set_run_font(sr, size=10.5, color=COLORS["muted"], italic=True)


def add_paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=COLORS["panel"])
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body, fill="EAF5FF", accent="4060C0"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=accent, size="10")
    cell = table.cell(0, 0)
    set_width(cell, 6.45)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.8, color=accent, bold=True)
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(0)
    br = bp.add_run(body)
    set_run_font(br, size=10, color=COLORS["ink"])


def add_screenshot_placeholder(doc, label, height_lines=4):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=COLORS["border"], size="10")
    cell = table.cell(0, 0)
    set_width(cell, 6.45)
    shade_cell(cell, COLORS["pale"])
    set_cell_margins(cell, top=260, bottom=260, start=160, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    set_run_font(r, size=10.5, color=COLORS["royal"], bold=True)
    for _ in range(max(0, height_lines - 1)):
        p.add_run("\n")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cr = cap.add_run("Area ini dapat diganti dengan screenshot asli dari game.")
    set_run_font(cr, size=8.5, color=COLORS["muted"], italic=True)


def add_table(doc, headers, rows, widths, header_fill="DCEEFF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table, color=COLORS["border"], size="8")

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_width(cell, widths[i])
        shade_cell(cell, header_fill)
        set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9.6, color=COLORS["panel"], bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_width(cell, widths[i])
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=COLORS["ink"])
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def build_cover(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    add_kicker(doc, "Buku Panduan")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("NetPro")
    set_run_font(r, size=42, color=COLORS["blue"], bold=True)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("Magang Jaringan")
    set_run_font(sr, size=22, color=COLORS["royal"], bold=True)

    if LOGO.exists():
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.paragraph_format.space_before = Pt(4)
        ip.paragraph_format.space_after = Pt(6)
        run = ip.add_run()
        run.add_picture(str(LOGO), width=Inches(3.9))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Game Edukasi Visual Novel untuk Teknik Komputer dan Jaringan")
    set_run_font(r, size=13, color=COLORS["muted"], italic=True)

    add_callout(
        doc,
        "Disusun oleh: Eko Muhammad Rizki",
        "Panduan ini menjelaskan tampilan, tombol, alur chapter, mini game, kuis refleksi, sistem skor, serta informasi pengembang dan sumber aset NetPro.",
        fill=COLORS["pale"],
        accent=COLORS["blue"],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    for name, hex_value in [
        ("Cyan", COLORS["cyan"]),
        ("Royal", COLORS["royal"]),
        ("Blue", COLORS["blue"]),
        ("Navy", COLORS["panel"]),
        ("Success", COLORS["success"]),
        ("Warning", COLORS["warning"]),
    ]:
        run = p.add_run(f" {name} #{hex_value} ")
        set_run_font(run, size=8.5, color=hex_value if name not in ("Warning", "Success", "Cyan") else COLORS["panel"], bold=True)

    page_break(doc)


def add_section(doc, n, title, subtitle, body=None, bullets=None, placeholder=None, callout=None):
    add_section_title(doc, n, title, subtitle)
    if placeholder:
        add_screenshot_placeholder(doc, placeholder)
    if body:
        for paragraph in body:
            add_paragraph(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)
    if callout:
        add_callout(doc, callout[0], callout[1], fill=callout[2], accent=callout[3])


def build_document():
    doc = Document()
    configure_document(doc)
    set_header_footer(doc.sections[0])

    build_cover(doc)

    add_section(
        doc,
        2,
        "Tentang Game",
        "Gambaran umum NetPro: Magang Jaringan.",
        body=[
            "NetPro: Magang Jaringan adalah game edukasi berbasis Visual Novel yang dirancang untuk siswa Teknik Komputer dan Jaringan (TKJ). Pemain berperan sebagai teknisi magang di PT. Nusanet Teknologi dan belajar memecahkan masalah jaringan komputer melalui cerita, dialog, pilihan interaktif, kuis refleksi, dan mini game crimping kabel.",
            "Materi yang dibawa meliputi pengabelan, subnetting, VLAN, routing, firewall, server, VoIP, QoS, VPN, dan disaster recovery. Format visual novel membuat materi teknis terasa dekat dengan situasi kerja magang.",
        ],
        placeholder="Screenshot: Tampilan pengenalan game / menu awal",
        callout=("Fokus Pembelajaran", "Belajar jaringan dari konteks lapangan: masalah muncul, pemain menganalisis, memilih solusi, lalu menerima umpan balik teknis.", COLORS["soft"], COLORS["royal"]),
    )
    page_break(doc)

    add_section(
        doc,
        3,
        "Tujuan dan Manfaat",
        "Alasan NetPro dibuat sebagai media belajar jaringan.",
        body=[
            "NetPro dibuat untuk membantu pemain memahami hubungan antara teori jaringan di kelas dan praktik teknisi di dunia industri. Pemain dapat belajar dari kesalahan tanpa risiko merusak perangkat atau mengganggu jaringan sungguhan.",
        ],
        bullets=[
            "Menjembatani teori dan praktik jaringan komputer.",
            "Memberi simulasi aman untuk mencoba konfigurasi dan troubleshooting.",
            "Mempersiapkan siswa TKJ sebelum magang atau prakerin.",
            "Mengubah materi jaringan yang berat menjadi pengalaman belajar yang naratif dan interaktif.",
        ],
        placeholder="Screenshot: Narasi awal program magang NetPro",
    )
    page_break(doc)

    add_section(
        doc,
        4,
        "Splash Screen dan Menu Utama",
        "Tampilan pertama saat game dibuka.",
        body=[
            "Ketika game dijalankan, pemain melihat intro atau splash video terlebih dahulu. Setelah itu, pemain diarahkan ke menu utama NetPro. Menu utama memakai tombol gambar dengan nuansa biru, ditempatkan di samping logo NetPro.",
            "Tombol utama yang tersedia adalah Mulai, Lanjutkan, Pilih Chapter, Pengaturan, Kredit, dan Keluar. Dari halaman ini pemain dapat memulai game baru, melanjutkan save, mengatur preferensi, membaca kredit, atau keluar dari aplikasi.",
        ],
        placeholder="Screenshot: Splash Screen dan Menu Utama NetPro",
    )
    page_break(doc)

    add_section_title(doc, 5, "Petunjuk Tombol", "Fungsi tombol dan navigasi utama.")
    add_screenshot_placeholder(doc, "Screenshot: Tombol menu utama dan quick menu", height_lines=3)
    add_table(
        doc,
        ["Tombol / Menu", "Fungsi"],
        [
            ("Mulai", "Memulai permainan baru dari awal."),
            ("Lanjutkan", "Membuka menu load untuk melanjutkan progress."),
            ("Pilih Chapter", "Menampilkan daftar chapter yang sudah terbuka."),
            ("Pengaturan", "Mengatur teks, audio, layar, dan preferensi Ren'Py."),
            ("Kredit", "Menampilkan informasi game, pengembang, dan teknologi."),
            ("Keluar", "Menutup game."),
            ("Simpan / Muat", "Menyimpan atau membuka data permainan."),
            ("Riwayat", "Melihat dialog yang sudah lewat."),
            ("Menu Utama", "Kembali ke menu utama dari dalam game."),
            ("Bantuan", "Melihat bantuan kontrol keyboard, mouse, atau gamepad."),
        ],
        [1.65, 4.75],
    )
    page_break(doc)

    add_section(
        doc,
        6,
        "Informasi Game, Kredit, dan Sumber",
        "Halaman tentang game dan kebutuhan pencantuman aset.",
        body=[
            "Halaman Kredit berisi identitas game, ringkasan konsep, pengembang, dan teknologi yang digunakan. Bagian sumber aset sebaiknya mencantumkan sumber gambar karakter, background, UI, musik, sound effect, movie intro, dan font.",
        ],
        bullets=[
            "Engine: Ren'Py Visual Novel Engine.",
            "Bahasa: Python dan Ren'Py Script.",
            "Platform target: Windows, Linux, macOS, dan Android.",
            "Versi game: 1.0.0.",
        ],
        placeholder="Screenshot: Halaman Kredit / Tentang Game",
    )
    page_break(doc)

    add_section_title(doc, 7, "Fitur Pilih Chapter", "Progress chapter dibuka bertahap.")
    add_screenshot_placeholder(doc, "Screenshot: Tampilan Pilih Chapter", height_lines=3)
    add_paragraph(doc, "Menu Pilih Chapter menampilkan daftar misi utama NetPro. Chapter 1 selalu terbuka, sedangkan chapter berikutnya akan terbuka setelah pemain menyelesaikan chapter sebelumnya.")
    add_numbered(
        doc,
        [
            "Chapter 1: Dasar Jaringan & Crimping Kabel.",
            "Chapter 2: VLAN & Segmentasi Jaringan.",
            "Chapter 3: WAN, NAT, & Security Firewall.",
            "Chapter 4: Server Administrator & VoIP.",
            "Chapter 5: Advanced Management & Disaster Recovery.",
        ],
    )
    add_callout(doc, "Catatan Progress", "Sistem pembuka chapter memakai persistent.chapter_unlocked, sehingga progress pemain dapat dipertahankan antar sesi.", COLORS["soft"], COLORS["royal"])
    page_break(doc)

    add_section(
        doc,
        8,
        "Alur Bermain Story Mode",
        "Cara pemain mengikuti cerita dan quest.",
        body=[
            "Pada awal permainan, pemain memasukkan nama karakter. Setelah itu, pemain mengikuti narasi magang sebagai teknisi jaringan baru. Cerita mempertemukan pemain dengan Pak Hendra sebagai mentor, Rafi sebagai sesama anak magang, Bu Dewi sebagai client, dan Pak Admin sebagai evaluator sistem.",
            "Saat quest muncul, pemain memilih jawaban atau aksi teknis. Jawaban benar membuat cerita lanjut dan skor bertambah. Jawaban salah menampilkan penjelasan atau hint, lalu pemain dapat mencoba kembali.",
        ],
        placeholder="Screenshot: Dialog karakter dan pilihan jawaban",
    )
    page_break(doc)

    add_section(
        doc,
        9,
        "Chapter 1: Dasar Jaringan dan Crimping",
        "Hari pertama magang, kabel UTP, OSI, APIPA, dan STP.",
        body=[
            "Chapter 1 memperkenalkan pemain pada dasar kerja teknisi jaringan. Pemain belajar memilih kabel untuk menghubungkan PC ke switch, melakukan crimping RJ45, memahami lapisan OSI, mengenali APIPA 169.254.x.x, dan menyelesaikan total network failure dengan konsep STP.",
        ],
        bullets=[
            "Quest kabel PC ke switch memberi pemahaman straight-through dan crossover.",
            "Mini game crimping melatih urutan warna T568B dan T568A.",
            "Bagian troubleshooting mengenalkan DHCP, APIPA, broadcast storm, dan Spanning Tree Protocol.",
            "Akhir chapter menampilkan skor, grade, dan rekap materi.",
        ],
        placeholder="Screenshot: Chapter 1 / Quest kabel dan jaringan",
    )
    page_break(doc)

    add_section_title(doc, 10, "Mini Game Crimping Kabel", "Menyusun 8 pin RJ45 sesuai standar.")
    add_screenshot_placeholder(doc, "Screenshot: Mini Game Crimping RJ45", height_lines=3)
    add_paragraph(doc, "Pemain memilih warna kabel untuk mengisi delapan slot pin RJ45. Setelah susunan lengkap, pemain menekan tombol cek untuk memvalidasi jawaban. Jika urutan benar, pemain mendapat poin dan lanjut. Jika salah, pemain dapat mengulang atau mereset susunan.")
    add_table(
        doc,
        ["Standar", "Urutan warna pin 1-8"],
        [
            ("T568B", "Putih Orange, Orange, Putih Hijau, Biru, Putih Biru, Hijau, Putih Cokelat, Cokelat."),
            ("T568A", "Putih Hijau, Hijau, Putih Orange, Biru, Putih Biru, Orange, Putih Cokelat, Cokelat."),
        ],
        [1.25, 5.15],
        header_fill=COLORS["pale"],
    )
    add_callout(doc, "Tombol Penting", "Cek Jawaban untuk validasi, Reset untuk menghapus susunan, dan Kembali untuk keluar dari mini game.", COLORS["soft"], COLORS["blue"])
    page_break(doc)

    add_section(
        doc,
        11,
        "Chapter 2: VLAN dan Segmentasi",
        "Subnetting, VLAN, inter-VLAN routing, dan port security.",
        body=[
            "Chapter 2 membawa pemain ke kasus penambahan divisi IT, Marketing, dan Finance. Jaringan yang semula flat perlu disegmentasi agar lebih aman dan efisien.",
        ],
        bullets=[
            "Subnetting 192.168.10.0/24 menjadi /26 untuk kebutuhan divisi.",
            "Konfigurasi VLAN 10, VLAN 20, dan VLAN 30 pada switch.",
            "Perintah switchport access vlan <ID> untuk memasukkan port ke VLAN tertentu.",
            "Inter-VLAN routing dengan router-on-a-stick.",
            "Port security violation shutdown untuk menghadapi perangkat asing.",
        ],
        placeholder="Screenshot: Chapter 2 / VLAN dan switch configuration",
    )
    page_break(doc)

    add_section(
        doc,
        12,
        "Chapter 3: WAN, NAT, dan Firewall",
        "Menghubungkan cabang, internet gateway, dan keamanan perimeter.",
        body=[
            "Chapter 3 fokus pada jaringan antar lokasi dan keamanan akses. Pemain menangani koneksi cabang ke HQ dengan dynamic routing OSPF, mengaktifkan NAT agar perangkat internal bisa mengakses internet, lalu merespons insiden firewall dengan ACL.",
        ],
        bullets=[
            "OSPF membantu routing antar jaringan cabang dan kantor pusat.",
            "NAT menerjemahkan alamat privat agar bisa keluar ke internet.",
            "Firewall dan ACL membatasi lalu lintas yang berisiko.",
        ],
        placeholder="Screenshot: Chapter 3 / WAN dan firewall incident",
    )
    page_break(doc)

    add_section(
        doc,
        13,
        "Chapter 4: Server Administrator dan VoIP",
        "Layanan server, komunikasi internal, dan filtering bandwidth.",
        body=[
            "Chapter 4 memperkenalkan peran teknisi jaringan dalam layanan server. Pemain mengatur DNS dan web server agar layanan dapat diakses melalui nama domain, menyiapkan VoIP/IP PBX untuk komunikasi kantor, serta mengatasi insiden bandwidth dengan web proxy atau filtering.",
        ],
        bullets=[
            "DNS memudahkan akses layanan memakai nama, bukan alamat IP.",
            "VoIP/IP PBX menyediakan nomor extension internal.",
            "Web proxy/filtering membantu mengendalikan penggunaan bandwidth.",
        ],
        placeholder="Screenshot: Chapter 4 / DNS, web server, dan VoIP",
    )
    page_break(doc)

    add_section(
        doc,
        14,
        "Chapter 5: Management dan Disaster Recovery",
        "QoS, VPN, backup, dan pemulihan insiden.",
        body=[
            "Final chapter menguji kesiapan pemain sebagai teknisi jaringan yang lebih matang. Pemain mengatur QoS agar aplikasi penting mendapat prioritas, membuat remote access VPN untuk pekerja jarak jauh, dan menangani insiden ransomware dengan disaster recovery protocol.",
        ],
        bullets=[
            "QoS menjaga aplikasi penting seperti meeting agar tidak terganggu.",
            "VPN memberi akses remote yang lebih aman.",
            "Disaster recovery menekankan backup, restore, dan pemulihan layanan.",
        ],
        placeholder="Screenshot: Chapter 5 / Disaster recovery dan VPN",
    )
    page_break(doc)

    add_section(
        doc,
        15,
        "Quiz dan Refleksi Materi",
        "Evaluasi singkat setelah chapter.",
        body=[
            "Setiap chapter memiliki kuis refleksi berisi lima pertanyaan. Kuis dipakai untuk memastikan pemain memahami materi sebelum melanjutkan ke chapter berikutnya.",
        ],
        bullets=[
            "Jawaban benar memberi +10 poin refleksi.",
            "Nilai sempurna memberi bonus +20 poin.",
            "Nilai baik memberi bonus +10 poin.",
            "Materi kuis mencakup kabel, OSI, APIPA, STP, VLAN, NAT, DNS, QoS, dan VPN.",
        ],
        placeholder="Screenshot: Tampilan quiz / refleksi materi",
    )
    page_break(doc)

    add_section(
        doc,
        16,
        "Sistem Skor, Grade, dan Progress",
        "Cara game menilai performa pemain.",
        body=[
            "Skor bertambah ketika pemain menyelesaikan quest utama, menjawab pilihan teknis dengan benar, menyelesaikan climax, dan meraih hasil baik pada kuis refleksi. Pada akhir chapter, game menampilkan evaluasi berupa skor dan grade.",
            "Progress chapter disimpan melalui sistem persistent chapter_unlocked. Pemain juga dapat memakai fitur simpan dan muat bawaan Ren'Py untuk melanjutkan sesi permainan.",
        ],
        placeholder="Screenshot: Rekap skor / hasil evaluasi chapter",
        callout=("Arah Evaluasi", "Skor tinggi memberi apresiasi performa teknisi. Skor sedang memberi saran latihan. Skor rendah mengarahkan pemain untuk mengulang dan memperkuat konsep.", COLORS["soft"], COLORS["royal"]),
    )
    page_break(doc)

    add_section_title(doc, 17, "Kondisi Benar, Salah, Restart, dan Exit", "Respon game terhadap tindakan pemain.")
    add_screenshot_placeholder(doc, "Screenshot: Feedback benar/salah dan pilihan ulang", height_lines=3)
    add_table(
        doc,
        ["Kondisi", "Penjelasan"],
        [
            ("Benar", "Cerita lanjut, sistem memberi umpan balik positif, dan skor bertambah."),
            ("Salah", "Game memberi penjelasan kesalahan atau hint, lalu pemain mencoba lagi."),
            ("Chapter selesai", "Rekap materi, skor, grade, dan pilihan lanjut ditampilkan."),
            ("Ulang / restart", "Pemain dapat mengulang mini game, kuis, atau chapter tertentu."),
            ("Exit / menu utama", "Pemain kembali ke menu utama atau keluar dari aplikasi."),
        ],
        [1.65, 4.75],
    )
    page_break(doc)

    add_section_title(doc, 18, "Profil Pengembang dan Penutup", "Informasi akhir guide book.")
    add_paragraph(doc, "Nama: Eko Muhammad Rizki")
    add_paragraph(doc, "Peran: Game Developer")
    add_paragraph(doc, "Proyek: Game edukasi jaringan berbasis Ren'Py untuk pembelajaran multimedia.")
    add_callout(
        doc,
        "Penutup",
        "NetPro: Magang Jaringan diharapkan membantu siswa TKJ memahami jaringan komputer secara interaktif, kontekstual, dan menyenangkan sebelum memasuki dunia kerja atau program magang.",
        COLORS["pale"],
        COLORS["blue"],
    )
    add_paragraph(doc, "Checklist sebelum dipublikasikan:")
    add_bullets(
        doc,
        [
            "Cover memuat judul, nama game, dan pengembang.",
            "Deskripsi game dan tujuan pembelajaran sudah jelas.",
            "Menu utama dan tombol sudah dijelaskan.",
            "Setiap fitur memiliki placeholder screenshot.",
            "Alur chapter, mini game, quiz, skor, dan progress dijelaskan.",
            "Profil pengembang dan sumber aset dicantumkan.",
        ],
    )

    for section in doc.sections:
        set_header_footer(section)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
